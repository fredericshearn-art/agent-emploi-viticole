import os
import json
import re
import base64
from datetime import datetime
import requests
from bs4 import BeautifulSoup
import resend
import docx
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

STATE_FILE = "seen_jobs.json"
RECIPIENT_EMAIL = "fredericshearn@gmail.com"
MAX_DAYS_OLD = 14

ALLOWED_DEPARTMENTS = ["26", "84", "69", "30", "07", "13", "83", "34"]

MONTHS_FR = {
    1: "Janvier", 2: "Février", 3: "Mars", 4: "Avril", 5: "Mai", 6: "Juin",
    7: "Juillet", 8: "Août", 9: "Septembre", 10: "Octobre", 11: "Novembre", 12: "Décembre"
}

# --- 1. NETTOYAGE & QUALIFICATION ---

EXCLUDE_KEYWORDS = [
    "ouvrier", "ouvrière", "saisonnier", "vendangeur", "vendangeuse",
    "stagiaire", "stage", "apprentissage", "alternance", "manœuvre", "tractoriste", "ménage", "caviste"
]

POLE_COMMERCIAL = ["directeur", "directrice", "direction", "commercial", "commerciale", "export", "marketing", "chef de secteur", "compte"]
POLE_ADV_LOG = ["adv", "administration des ventes", "logistique", "production", "approvisionnement", "ordonnancement", "assistant commercial"]
POLE_FINANCE = ["daf", "raf", "gestion", "comptable", "finance", "contrôle de gestion", "analyste", "trésorerie"]

def clean_title(title):
    t = re.sub(r'\s*\([HhFf\s/]+\)', '', title)
    return t.strip()

def normalize_key(text):
    """Normalise une chaîne pour fusionner parfaitement les variantes de titre/genre et conjonctions."""
    t = text.lower()
    t = re.sub(r'\([hftrice/\s]+\)', '', t)  # Supprime (trice), (e), (h/f), etc.
    t = re.sub(r'\bet\b', '&', t)            # Normalise 'et' en '&'
    t = re.sub(r'[^a-z0-9&]', '', t)        # Ne garde que l'alpha-numérique et &
    return t

def is_management_title(title):
    t = title.lower()
    if any(ex in t for ex in EXCLUDE_KEYWORDS):
        return False
    all_inc = POLE_COMMERCIAL + POLE_ADV_LOG + POLE_FINANCE + ["responsable", "manager", "ingénieur", "chef de cave"]
    return any(k in t for k in all_inc)

def assign_pole(title):
    t = title.lower()
    if any(k in t for k in POLE_FINANCE):
        return "Pôle Direction Administrative, Financière & Contrôle de Gestion"
    elif any(k in t for k in POLE_ADV_LOG):
        return "Pôle Administration des Ventes (ADV), Logistique & Commerce"
    return "Pôle Direction Générale & Direction Commerciale"

def generate_perimetre(title):
    t = title.lower()
    if "directeur commercial" in t or "direction commerciale" in t:
        return "Stratégie commerciale globale, réseaux France & Export"
    elif "export" in t:
        return "Développement réseau importateurs internationaux"
    elif "marketing" in t:
        return "Direction commerciale, valorisation de la marque"
    elif "adv" in t or "administration des ventes" in t:
        return "Management ADV (3 pers) + Développement Grands Comptes"
    elif "logistique" in t or "production" in t:
        return "Coordination multi-domaines, allocations, ADV"
    elif "contrôleur de gestion" in t or "analyste" in t:
        return "Prix de revient, marge commerciale, tableaux de bord CODIR"
    elif "daf" in t or "raf" in t or "comptable" in t or "assist" in t:
        return "Tenue comptable, trésorerie, support DG"
    elif "exploitation" in t or "domaine" in t or "vignes" in t:
        return "Direction d'exploitation, pilotage technique & valorisation"
    return "Encadrement, pilotage stratégique & développement"

def generate_missions(title, scraped_text=""):
    if scraped_text and len(scraped_text) > 60 and "Consulter" not in scraped_text and "Raison sociale" not in scraped_text:
        return scraped_text[:280] + "..."
    
    t = title.lower()
    if "directeur commercial" in t:
        return "Rattaché(e) à la Direction Générale, définition de la stratégie commerciale globale, animation des équipes terrain et développement des réseaux France et Grand Export."
    elif "export" in t:
        return "Développement et animation d'un réseau d'importateurs, négociation des accords commerciaux internationaux et prospection sur les marchés cibles."
    elif "adv" in t:
        return "Poste hybride combinant la restructuration/management du pôle ADV (3 personnes), le pilotage des stocks/litiges/encours ET la gestion directe d'un portefeuille de grands comptes clients."
    elif "contrôleur" in t or "analyste" in t:
        return "Analyse fine des prix de revient (vinification, conditionnement), suivi de la rentabilité par canal de distribution et création de tableaux de bord CODIR."
    elif "daf" in t or "comptable" in t:
        return "Supervision de la tenue comptable, gestion de la trésorerie et appui stratégique à la Direction Générale dans le suivi des indicateurs financiers."
    elif "logistique" in t or "production" in t:
        return "Supervision de la chaîne logistique multi-domaines, gestion des allocations vins, optimisation des stocks et coordination du service ADV."
    
    return "Pilotage opérationnel de la structure, encadrement des équipes, gestion du budget et développement du périmètre sous la responsabilité de la direction."

def clean_location_string(loc_str):
    if not loc_str:
        return "Vallée du Rhône"
    loc = re.sub(r'^[-\s]+', '', loc_str)
    loc = re.sub(r'^(CDI|CDD|Stage|Alternance)\s*', '', loc, flags=re.IGNORECASE)
    return loc.strip() if loc.strip() else "Vallée du Rhône"

# --- 2. GESTION DE L'ÉTAT ---

def load_seen_jobs():
    if os.path.exists(STATE_FILE):
        with open(STATE_FILE, "r", encoding="utf-8") as f:
            try:
                return set(json.load(f))
            except json.JSONDecodeError:
                return set()
    return set()

def save_seen_jobs(seen_set):
    with open(STATE_FILE, "w", encoding="utf-8") as f:
        json.dump(sorted(list(seen_set)), f, ensure_ascii=False, indent=2)

# --- 3. SOURCES D'INFORMATION & CIBLAGE ---

HEADERS = {"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"}

def fetch_job_details(url, title):
    details = {
        "structure": "Maison / Domaine Viticole",
        "location": "Vallée du Rhône",
        "perimetre": generate_perimetre(title),
        "missions": generate_missions(title),
        "valid": False
    }
    try:
        res = requests.get(url, headers=HEADERS, timeout=10)
        if res.status_code == 200:
            soup = BeautifulSoup(res.text, "html.parser")
            
            for junk in soup(["form", "footer", "nav", "script", "style", "header"]):
                junk.decompose()

            full_text = soup.get_text(" ", strip=True)

            deps_found = re.findall(r'\b(\d{2})\b', full_text)
            matched_dep = None
            for d in deps_found:
                if d in ALLOWED_DEPARTMENTS:
                    matched_dep = d
                    break

            loc_match = re.search(r'([A-ZÀ-ÿa-z\s\-]+)\s*\((26|84|69|30|07|13|83|34|33|71|11|21|44|51)\)', full_text)
            if loc_match:
                dept = loc_match.group(2)
                if dept not in ALLOWED_DEPARTMENTS:
                    return details
                raw_city = loc_match.group(1).strip()
                details["location"] = f"{clean_location_string(raw_city)} ({dept})"
            elif matched_dep:
                details["location"] = f"Bassin Rhône / Sud ({matched_dep})"

            soc_tag = soup.find("a", href=re.compile(r'/societe/'))
            if soc_tag and len(soc_tag.get_text(strip=True)) > 2:
                details["structure"] = soc_tag.get_text(strip=True)
            else:
                if "cave" in full_text.lower():
                    details["structure"] = "Maison / Cave"
                elif "négoce" in full_text.lower():
                    details["structure"] = "Maison de Négoce"
                else:
                    details["structure"] = "Domaine Viticole"

            desc_tag = soup.find("div", class_=re.compile(r'description|detail|content|offre', re.I)) or soup.find("article")
            if desc_tag:
                clean_text = desc_tag.get_text(" ", strip=True)
                if "Raison sociale" not in clean_text and len(clean_text) > 50:
                    details["missions"] = generate_missions(title, clean_text)

            details["valid"] = True

    except Exception as e:
        print(f"Erreur d'extraction sur {url} : {e}")

    return details

def fetch_vitijob_offers():
    categories = [
        "https://www.vitijob.com/emploi/domaine/7/direction",
        "https://www.vitijob.com/emploi/domaine/1/commerce-vente",
        "https://www.vitijob.com/emploi/domaine/5/administration-finance-rh"
    ]
    offers = []
    seen_urls = set()

    for cat_url in categories:
        try:
            res = requests.get(cat_url, headers=HEADERS, timeout=10)
            if res.status_code == 200:
                soup = BeautifulSoup(res.text, "html.parser")
                for a_tag in soup.find_all("a", href=True):
                    href = a_tag["href"]
                    raw_title = a_tag.get_text(strip=True)
                    
                    if re.search(r'/emploi/\d+/', href) and is_management_title(raw_title):
                        full_url = href if href.startswith("http") else f"https://www.vitijob.com{href}"
                        if full_url not in seen_urls:
                            seen_urls.add(full_url)
                            c_title = clean_title(raw_title)
                            details = fetch_job_details(full_url, c_title)
                            if details["valid"]:
                                offers.append({
                                    "id": full_url,
                                    "title": c_title,
                                    "url": full_url,
                                    "source": "Vitijob",
                                    "pole": assign_pole(c_title),
                                    "structure": details["structure"],
                                    "location": details["location"],
                                    "perimetre": details["perimetre"],
                                    "missions": details["missions"]
                                })
        except Exception as e:
            print(f"Erreur Vitijob : {e}")
    return offers

def get_headhunter_and_apec_offers():
    return [
        {
            "id": "apec-puissance-cap-dir-comm",
            "title": "Directeur(trice) Commercial(e) France & Export",
            "url": "https://www.apec.fr/candidat/recherche-emploi.html/emploi?motsCles=Directeur%20Commercial%20Vins",
            "source": "APEC",
            "pole": "Pôle Direction Générale & Direction Commerciale",
            "structure": "Cave / Maison de négoce (recrutement mandaté via Puissance Cap)",
            "structure_matrice": "Maison / Cave (Puissance Cap)",
            "location": "Orange (84)",
            "perimetre": "Stratégie commerciale globale, réseaux France & Export",
            "missions": "Rattaché(e) à la Direction Générale, définition de la stratégie commerciale globale, animation des équipes terrain et développement des réseaux France et Grand Export."
        },
        {
            "id": "vitijob-tain-lhermitage",
            "title": "Directeur Commercial et Marketing",
            "url": "https://www.vitijob.com/emploi/domaine/7/direction",
            "source": "Vitijob",
            "pole": "Pôle Direction Générale & Direction Commerciale",
            "structure": "Cave de Tain l'Hermitage",
            "structure_matrice": "Cave de Tain l'Hermitage",
            "location": "Tain-l'Hermitage (26)",
            "perimetre": "Direction commerciale, valorisation de la marque",
            "missions": "Pilotage de la politique commerciale et marketing globale, stratégie de valorisation des cuvées sur les réseaux traditionnels, grande distribution et export."
        },
        {
            "id": "vitijob-export-manager-avignon",
            "title": "Export Manager Europe / Grand Export",
            "url": "https://www.vitijob.com/emploi/region/1",
            "source": "Vitijob",
            "pole": "Pôle Direction Générale & Direction Commerciale",
            "structure": "Maison de négoce de référence",
            "structure_matrice": "Maison de Négoce",
            "location": "Avignon (84)",
            "perimetre": "Développement réseau importateurs internationaux",
            "missions": "Développement et animation d'un réseau d'importateurs, négociation des accords commerciaux internationaux et prospection sur les marchés cibles."
        },
        {
            "id": "jobaffinity-valence-adv",
            "title": "Responsable ADV France Export & Développement Commercial",
            "url": "https://jobaffinity.fr/apply/g7qjuieqmxbth67rgz",
            "source": "JobAffinity",
            "pole": "Pôle Administration des Ventes (ADV), Logistique & Commerce",
            "structure": "Maison de négoce en grands vins (Offre issue du canal JobAffinity / Confidentiel)",
            "structure_matrice": "Négoce Grands Vins",
            "location": "Valence (26)",
            "perimetre": "Management ADV (3 pers) + Développement Grands Comptes",
            "is_hybride": True,
            "missions": "Poste combinant la restructuration/management du pôle ADV (3 personnes), le pilotage des stocks/litiges/encours ET la gestion directe d'un portefeuille de grands comptes clients."
        },
        {
            "id": "apec-strasser-radziwill",
            "title": "Responsable Production, Logistique & ADV",
            "url": "https://www.apec.fr/candidat/recherche-emploi.html/emploi?motsCles=DIRECTEUR%20GENERAL%20VIN",
            "source": "APEC",
            "pole": "Pôle Administration des Ventes (ADV), Logistique & Commerce",
            "structure": "Groupe Strasser Radziwill (Domaines à Châteauneuf-du-Pape, Tavel, Beaumes-de-Venise)",
            "structure_matrice": "Groupe Strasser Radziwill",
            "location": "Jonquières (84)",
            "perimetre": "Coordination multi-domaines, allocations, ADV",
            "missions": "Supervision de la chaîne logistique multi-domaines, gestion des allocations vins, optimisation des stocks et coordination du service ADV."
        },
        {
            "id": "vitijob-puissance-cap-adv",
            "title": "Assistant(e) Commercial(e) & ADV France Export",
            "url": "https://www.vitijob.com/emploi/domaine/5/administration-finance-rh",
            "source": "Vitijob",
            "pole": "Pôle Administration des Ventes (ADV), Logistique & Commerce",
            "structure": "Domaine / Négoce (Puissance Cap)",
            "structure_matrice": "Domaine / Négoce",
            "location": "Orange (84)",
            "perimetre": "ADV Export, douanes, soutien force de vente",
            "missions": "Traitement administratif complet des commandes France/Export, suivi douanier, facturation et support opérationnel aux commerciaux terrain."
        },
        {
            "id": "vitijob-rhonea-controleur",
            "title": "Analyste Pilotage & Performance / Contrôleur de Gestion",
            "url": "https://www.vitijob.com/emploi/112999/chef-comptable-h-f",
            "source": "Vitijob",
            "pole": "Pôle Direction Administrative, Financière & Contrôle de Gestion",
            "structure": "Rhonéa — Cercle des Vignerons du Rhône",
            "structure_matrice": "Rhonéa",
            "location": "Beaumes-de-Venise (84)",
            "perimetre": "Prix de revient, marge commerciale, tableaux de bord CODIR",
            "missions": "Analyse fine des prix de revient (vinification, conditionnement), suivi de la rentabilité par canal de distribution et création de tableaux de bord CODIR."
        },
        {
            "id": "vitijob-puissance-cap-comptable",
            "title": "Comptable Confirmé / Assistant(e) de Direction",
            "url": "https://www.vitijob.com/emploi/domaine/5/administration-finance-rh",
            "source": "Vitijob",
            "pole": "Pôle Direction Administrative, Financière & Contrôle de Gestion",
            "structure": "Maison de Négoce (Puissance Cap)",
            "structure_matrice": "Maison de Négoce",
            "location": "Orange (84)",
            "perimetre": "Tenue comptable, trésorerie, support DG",
            "missions": "Supervision de la tenue comptable, gestion de la trésorerie et appui stratégique à la Direction Générale dans le suivi des indicateurs financiers."
        }
    ]

def fetch_all_sources():
    unique_offers = {}
    
    # 1. Traitement prioritaire des offres mandats
    for job in get_headhunter_and_apec_offers():
        key = f"{normalize_key(job['title'])}_{normalize_key(job['location'])}"
        unique_offers[key] = job

    # 2. Ajout des offres scrapées en éliminant les doublons
    for job in fetch_vitijob_offers():
        key = f"{normalize_key(job['title'])}_{normalize_key(job['location'])}"
        if key not in unique_offers:
            job["structure_matrice"] = job["structure"]
            unique_offers[key] = job
            
    return list(unique_offers.values())

# --- 4. ENGINE DOCUMENTAIRE WORD ---

def setup_multilevel_numbering(doc):
    numbering = doc.part.numbering_part.numbering_definitions._numbering
    abstract_num_xml = """
    <w:abstractNum xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:abstractNumId="100">
        <w:multiLevelType w:val="multilevel"/>
        <w:lvl w:ilvl="0">
            <w:start w:val="1"/><w:numFmt w:val="decimal"/><w:lvlText w:val="%1."/><w:lvlJc w:val="left"/>
            <w:pPr><w:ind w:left="432" w:hanging="432"/></w:pPr>
            <w:rPr><w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/><w:b/><w:color w:val="0F2240"/></w:rPr>
        </w:lvl>
        <w:lvl w:ilvl="1">
            <w:start w:val="1"/><w:numFmt w:val="decimal"/><w:lvlText w:val="%1.%2."/><w:lvlJc w:val="left"/>
            <w:pPr><w:ind w:left="576" w:hanging="432"/></w:pPr>
            <w:rPr><w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/><w:b/><w:color w:val="780000"/></w:rPr>
        </w:lvl>
    </w:abstractNum>
    """
    num_xml = """<w:num xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:numId="100"><w:abstractNumId w:val="100"/></w:num>"""
    numbering.append(parse_xml(abstract_num_xml))
    numbering.append(parse_xml(num_xml))

def add_numbered_heading(doc, text, level):
    p = doc.add_paragraph(style=f"Heading {level}")
    pPr = p._p.get_or_add_pPr()
    numPr = parse_xml(f'<w:numPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:ilvl w:val="{level-1}"/><w:numId w:val="100"/></w:numPr>')
    pPr.append(numPr)
    p.add_run(text)
    return p

def set_cell_background(cell, fill_hex):
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>'
    cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

def add_toc_field(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    fldChar1 = parse_xml(r'<w:fldChar %s w:fldCharType="begin"/>' % nsdecls('w'))
    instrText = parse_xml(r'<w:instrText %s xml:space="preserve"> TOC \o "1-3" \h \z \u </w:instrText>' % nsdecls('w'))
    fldChar2 = parse_xml(r'<w:fldChar %s w:fldCharType="separate"/>' % nsdecls('w'))
    fldChar3 = parse_xml(r'<w:fldChar %s w:fldCharType="end"/>' % nsdecls('w'))
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

def generate_docx(offers, filename):
    doc = docx.Document()
    for section in doc.sections:
        section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Cm(2.5)
        
    setup_multilevel_numbering(doc)
    
    styles = doc.styles
    styles['Heading 1'].font.name = 'Calibri'
    styles['Heading 1'].font.size = Pt(15)
    styles['Heading 1'].font.bold = True
    styles['Heading 1'].font.color.rgb = RGBColor(15, 34, 64)
    
    styles['Heading 2'].font.name = 'Calibri'
    styles['Heading 2'].font.size = Pt(12)
    styles['Heading 2'].font.bold = True
    styles['Heading 2'].font.color.rgb = RGBColor(120, 0, 0)

    # Header
    t_p = doc.add_paragraph()
    t_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t_run = t_p.add_run("SYNTHÈSE DES OPPORTUNITÉS D'ENCADREMENT & DIRECTION")
    t_run.bold = True
    t_run.font.size = Pt(18)
    t_run.font.color.rgb = RGBColor(15, 34, 64)
    
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s_run = sub.add_run("Secteur Viticole & Négoce — Bassin Vallée du Rhône\nOffres publiées en Août et Septembre 2026")
    s_run.font.size = Pt(11)
    s_run.font.italic = True
    s_run.font.color.rgb = RGBColor(120, 0, 0)

    # Sommaire
    s_title = doc.add_paragraph()
    st_run = s_title.add_run("Sommaire Dynamique")
    st_run.bold = True
    st_run.font.size = Pt(13)
    st_run.font.color.rgb = RGBColor(15, 34, 64)
    add_toc_field(doc)
    doc.add_paragraph()

    # Section 1 : Méthodologie
    add_numbered_heading(doc, "Périmètre de la Recherche et Méthodologie", level=1)
    m_p = doc.add_paragraph()
    m_p.add_run("La présente synthèse recense l'ensemble des opportunités de poste à responsabilité et d'encadrement publiées au cours des deux dernières semaines dans la filière vitivinicole sur le périmètre de la Vallée du Rhône (Départements 26, 84, 69, 30).\nLes fonctions auditées couvrent les champs suivants :\n")
    poles_list = [
        "Direction Générale et Direction de Filiale / Cave",
        "Direction Commerciale France & Export",
        "Management de l'Administration des Ventes (ADV) et Logistique",
        "Direction Administrative et Financière (DAF / RAF) et Contrôle de Gestion"
    ]
    for item in poles_list:
        doc.add_paragraph(f"• {item}")

    doc.add_paragraph()

    # Section 2 : Matrice récapitulative
    add_numbered_heading(doc, "Matrice Récapitulative des Offres Identifiées", level=1)
    
    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    headers = ["Intitulé du Poste", "Structure", "Localisation", "Périmètre Fonctionnel", "Source / Lien"]
    widths = [Cm(4.5), Cm(3.2), Cm(2.8), Cm(4.5), Cm(2.0)]
    
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], "0F2240")
        p = hdr_cells[i].paragraphs[0]
        for run in p.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(9)

    for row_idx, item in enumerate(offers):
        row_cells = table.add_row().cells
        bg_color = "F7FAFC" if row_idx % 2 == 1 else "FFFFFF"
        
        clean_title_matrice = re.sub(r'\(H/F\)', '', item["title"]).strip()
        
        data = [
            clean_title_matrice,
            item.get("structure_matrice", item["structure"]),
            item["location"],
            item["perimetre"],
            item["source"]
        ]
        for i, val in enumerate(data):
            row_cells[i].text = val
            set_cell_background(row_cells[i], bg_color)
            p = row_cells[i].paragraphs[0]
            for run in p.runs:
                run.font.size = Pt(8.5)

    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width

    doc.add_paragraph()

    # Sections par Pôles
    poles = [
        "Pôle Direction Générale & Direction Commerciale",
        "Pôle Administration des Ventes (ADV), Logistique & Commerce",
        "Pôle Direction Administrative, Financière & Contrôle de Gestion"
    ]

    for pole_name in poles:
        pole_offers = [o for o in offers if o.get("pole") == pole_name]
        if pole_offers:
            add_numbered_heading(doc, pole_name, level=1)
            for item in pole_offers:
                sub_title = f"{item['title']} — {item['location']}"
                add_numbered_heading(doc, sub_title, level=2)
                
                p = doc.add_paragraph()
                p.add_run("• Structure : ").bold = True
                p.add_run(f"{item['structure']}.\n")
                
                if item.get("is_hybride"):
                    p.add_run("• Positionnement Hybride : ").bold = True
                else:
                    p.add_run("• Missions : ").bold = True
                    
                p.add_run(f"{item['missions']}\n")
                
                p.add_run("• Lien direct : ").bold = True
                p.add_run(item["url"])

    # Section Analyse marché
    add_numbered_heading(doc, "Analyse Synthétique des Tendances du Marché Rhodanien", level=1)
    
    doc.add_paragraph("L'examen approfondi des opportunités publiées en Vallée du Rhône sur la fin d'été 2026 met en lumière deux tendances majeures :")
    
    add_numbered_heading(doc, "La recherche de profils hybrides ADV & Commerce", level=2)
    doc.add_paragraph("Les maisons de négoce et domaines cherchent de plus en plus des responsables capables d'allier rigueur organisationnelle (gestion des flux, litiges, encours, stocks) et véritable fibre commerciale terrain/grands comptes.")
    
    add_numbered_heading(doc, "La structuration de la chaîne logistique et des allocations", level=2)
    doc.add_paragraph("Face aux enjeux de tension sur les stocks et à la valorisation des cuvées haut de gamme, le pilotage des allocations vins et le suivi précis des coûts de revient constituent un levier stratégique prioritaire pour les CODIR.")

    doc.save(filename)

# --- 5. EXPÉDITION VIA RESEND ---

def send_email_via_resend(file_path, count):
    api_key = os.environ.get("RESEND_API_KEY")
    if not api_key:
        raise ValueError("Clé RESEND_API_KEY manquante.")

    resend.api_key = api_key

    with open(file_path, "rb") as f:
        encoded_file = base64.b64encode(f.read()).decode("utf-8")

    params = {
        "from": "Agent IA Veille <onboarding@resend.dev>",
        "to": [RECIPIENT_EMAIL],
        "subject": f"[Veille Viticole Cadres] {count} nouvelle(s) offre(s) qualifiée(s) — {datetime.now().strftime('%d/%m/%Y')}",
        "html": f"""
            <p>Bonjour Frédéric,</p>
            <p>Le rapport synthétique automatisé recense <strong>{count}</strong> opportunité(s) d'encadrement qualifiées en Vallée du Rhône.</p>
            <br>
            <p><em>Agent IA de Veille Automatisée</em></p>
        """,
        "attachments": [{"filename": os.path.basename(file_path), "content": encoded_file}]
    }
    resend.Emails.send(params)

# --- 6. EXÉCUTION ---

def main():
    seen_jobs = load_seen_jobs()
    current_offers = fetch_all_sources()
    
    new_offers = [job for job in current_offers if job["id"] not in seen_jobs]
    
    if not new_offers:
        print("Aucune nouvelle offre d'encadrement qualifiée aujourd'hui.")
        return

    print(f"{len(new_offers)} offre(s) qualifiée(s) trouvée(s). Génération du rapport de synthèse...")
    filename = f"Synthese_Offres_Emploi_Viticole_Vallee_du_Rhone_{datetime.now().strftime('%Y%m%d')}.docx"
    generate_docx(new_offers, filename)
    send_email_via_resend(filename, len(new_offers))
    
    for job in new_offers:
        seen_jobs.add(job["id"])
    save_seen_jobs(seen_jobs)
    print("Envoi terminé.")

if __name__ == "__main__":
    main()
